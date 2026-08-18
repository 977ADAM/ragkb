"""Загрузка документов разных форматов в единое внутреннее представление.

Каждый загрузчик возвращает список Block — абзац/заголовок/строку таблицы —
с метаданными о происхождении. Блоки, а не сырой текст, позволяют дальше
резать документ по смысловым границам и сохранять ссылку на страницу.
"""
from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt", ".html", ".htm"}


@dataclass
class Block:
    """Смысловой блок исходного документа."""
    text: str
    kind: str = "paragraph"          # paragraph | heading | table | list
    level: int = 0                   # уровень заголовка (1..6), 0 для остального
    page: int | None = None          # номер страницы для PDF
    meta: dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    doc_id: str
    path: str
    title: str
    blocks: list[Block]
    checksum: str
    meta: dict[str, Any] = field(default_factory=dict)


class UnsupportedFormat(Exception):
    pass


def discover(root: str | Path) -> list[Path]:
    """Рекурсивно находит все поддерживаемые файлы, игнорируя служебные."""
    root = Path(root)
    if root.is_file():
        return [root] if root.suffix.lower() in SUPPORTED_EXTENSIONS else []
    files = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith(".") or "~$" in path.name:
            continue
        if path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(path)
    return files


def load(path: str | Path) -> Document:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        blocks, meta = _load_pdf(path)
    elif suffix == ".docx":
        blocks, meta = _load_docx(path)
    elif suffix in {".md", ".markdown"}:
        blocks, meta = _load_markdown(path)
    elif suffix in {".html", ".htm"}:
        blocks, meta = _load_html(path)
    elif suffix == ".txt":
        blocks, meta = _load_text(path)
    else:
        raise UnsupportedFormat(f"Формат не поддерживается: {path.suffix}")

    blocks = [b for b in blocks if b.text.strip()]
    checksum = hashlib.sha256(path.read_bytes()).hexdigest()[:16]
    title = meta.pop("title", None) or _guess_title(blocks) or path.stem
    return Document(
        doc_id=hashlib.sha1(str(path.resolve()).encode()).hexdigest()[:12],
        path=str(path),
        title=title,
        blocks=blocks,
        checksum=checksum,
        meta={"format": suffix.lstrip("."), **meta},
    )


# --------------------------------------------------------------------------- PDF

def _load_pdf(path: Path) -> tuple[list[Block], dict[str, Any]]:
    text_by_page = _pdf_pages(path)
    blocks: list[Block] = []
    for page_no, page_text in enumerate(text_by_page, start=1):
        page_text = _fix_hyphenation(page_text)
        for para in _pdf_paragraphs(page_text):
            kind, level = _classify_line(para)
            blocks.append(Block(text=para, kind=kind, level=level, page=page_no))
    return blocks, {"pages": len(text_by_page)}


def _pdf_paragraphs(page_text: str) -> list[str]:
    """Собирает строки страницы в абзацы.

    В PDF нет понятия абзаца — извлекатель отдаёт поток строк. Пустая строка
    встречается не всегда, поэтому дополнительно склеиваем строки до тех пор,
    пока предыдущая не закончилась знаком конца предложения, и отделяем
    строки, похожие на заголовок.
    """
    paragraphs: list[str] = []
    current: list[str] = []

    def flush() -> None:
        if current:
            paragraphs.append(re.sub(r"\s+", " ", " ".join(current)).strip())
            current.clear()

    for raw_line in page_text.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        kind, _ = _classify_line(line)
        if kind == "heading":
            flush()
            paragraphs.append(line)
            continue
        current.append(line)
        # Конец предложения на границе строки — вероятная граница абзаца.
        if line.endswith((".", "!", "?", ":", ";")):
            flush()
    flush()
    return [p for p in paragraphs if p]


def _pdf_pages(path: Path) -> list[str]:
    """Извлекает текст постранично. pdfplumber точнее, pypdf — запасной вариант."""
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(str(path)) as pdf:
            return [(page.extract_text() or "") for page in pdf.pages]
    except ImportError:
        pass
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedFormat(
            "Для чтения PDF установите pypdf или pdfplumber"
        ) from exc
    reader = PdfReader(str(path))
    return [(page.extract_text() or "") for page in reader.pages]


def _fix_hyphenation(text: str) -> str:
    """PDF часто разрывает слова переносом на границе строки — склеиваем обратно."""
    return re.sub(r"(\w)-\n(\w)", r"\1\2", text)


# -------------------------------------------------------------------------- DOCX

def _load_docx(path: Path) -> tuple[list[Block], dict[str, Any]]:
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedFormat("Для чтения DOCX установите python-docx") from exc

    document = docx.Document(str(path))
    blocks: list[Block] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading") or style.startswith("заголовок"):
            level = _int_or(style.split()[-1], 1)
            blocks.append(Block(text=text, kind="heading", level=level))
        else:
            blocks.append(Block(text=text, kind="paragraph"))

    # Таблицы: каждая строка становится отдельным блоком вида "колонка: значение".
    for t_idx, table in enumerate(document.tables):
        header = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            pairs = [f"{h}: {v}" for h, v in zip(header, cells) if v]
            if pairs:
                blocks.append(
                    Block(text="; ".join(pairs), kind="table", meta={"table": t_idx})
                )

    core = document.core_properties
    meta = {"title": (core.title or "").strip() or None, "author": core.author or ""}
    return blocks, {k: v for k, v in meta.items() if v}


# ---------------------------------------------------------------------- Markdown

def _load_markdown(path: Path) -> tuple[list[Block], dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    blocks: list[Block] = []
    in_code = False
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            joined = "\n".join(buffer).strip()
            if joined:
                blocks.append(Block(text=joined, kind="paragraph"))
            buffer.clear()

    for line in text.splitlines():
        if line.strip().startswith("```"):
            in_code = not in_code
            buffer.append(line)
            if not in_code:
                flush()
            continue
        if in_code:
            buffer.append(line)
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            flush()
            blocks.append(
                Block(text=heading.group(2).strip(), kind="heading", level=len(heading.group(1)))
            )
            continue
        if not line.strip():
            flush()
            continue
        buffer.append(line)
    flush()
    return blocks, {}


# -------------------------------------------------------------------------- HTML

def _load_html(path: Path) -> tuple[list[Block], dict[str, Any]]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        return _load_text_from_string(_strip_tags(raw)), {}

    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    blocks: list[Block] = []
    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td"]):
        text = el.get_text(" ", strip=True)
        if not text:
            continue
        if el.name.startswith("h"):
            blocks.append(Block(text=text, kind="heading", level=int(el.name[1])))
        elif el.name == "li":
            blocks.append(Block(text=text, kind="list"))
        else:
            blocks.append(Block(text=text, kind="paragraph"))
    title = soup.title.get_text(strip=True) if soup.title else None
    return blocks, {"title": title} if title else {}


def _strip_tags(html: str) -> str:
    return re.sub(r"<[^>]+>", " ", html)


# -------------------------------------------------------------------------- TXT

def _load_text(path: Path) -> tuple[list[Block], dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return _load_text_from_string(text), {}


def _load_text_from_string(text: str) -> list[Block]:
    blocks = []
    for para in _split_paragraphs(text):
        kind, level = _classify_line(para)
        blocks.append(Block(text=para, kind=kind, level=level))
    return blocks


# ------------------------------------------------------------------- утилиты

def _split_paragraphs(text: str) -> Iterable[str]:
    for chunk in re.split(r"\n\s*\n", text):
        cleaned = re.sub(r"[ \t]+", " ", chunk).strip()
        if cleaned:
            yield cleaned


def _classify_line(text: str) -> tuple[str, int]:
    """Эвристика для plain-text: короткая строка без точки на конце — вероятно заголовок."""
    stripped = text.strip()
    if len(stripped) > 90 or "\n" in stripped:
        return "paragraph", 0
    if re.match(r"^\d+(\.\d+)*\.?\s+\S", stripped) and not stripped.endswith("."):
        return "heading", stripped.count(".") + 1
    if stripped.isupper() and len(stripped) > 3:
        return "heading", 1
    return "paragraph", 0


def _guess_title(blocks: list[Block]) -> str | None:
    for block in blocks[:5]:
        if block.kind == "heading":
            return block.text
    return blocks[0].text[:80] if blocks else None


def _int_or(value: str, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default
